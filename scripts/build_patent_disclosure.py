#!/usr/bin/env python3
"""Build the DVCL patent disclosure from the supplied Word template."""

from __future__ import annotations

import argparse
import math
import tempfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PATENT_TITLE = "一种基于双视图对比学习的鲁棒异构图节点分类方法、系统、设备及存储介质"
SHORT_TITLE = "DVCL鲁棒异构图节点分类方法及系统"

BODY_FONT_CN = "宋体"
BODY_FONT_LATIN = "Times New Roman"
HEADING_FONT_CN = "黑体"
MATH_FONT = "Cambria Math"


def set_run_font(run, chinese=BODY_FONT_CN, latin=BODY_FONT_LATIN, size=12, bold=False, color=None):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), chinese)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_body(paragraph, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Cm(0.74) if first_line else Cm(0)
    fmt.widow_control = True
    for run in paragraph.runs:
        set_run_font(run, size=size)


def add_body(doc, text, first_line=True, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    set_paragraph_body(p, first_line=first_line)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.widow_control = True
    if level == 1:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
        size = 16
    elif level == 2:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        size = 14
    else:
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(3)
        size = 12
    r = p.add_run(text)
    set_run_font(r, chinese=HEADING_FONT_CN, latin="Arial", size=size, bold=True)
    return p


def add_center_title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(16)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, chinese=HEADING_FONT_CN, latin="Arial", size=size, bold=True)
    return p


def add_numbered(doc, number, text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.74 * level)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{number} ")
    set_run_font(r1, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("• " + text)
    set_run_font(r)
    return p


def add_formula(doc, formula, number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_together = True
    r = p.add_run(f"{formula}    （{number}）")
    set_run_font(r, chinese=MATH_FONT, latin=MATH_FONT, size=11.5)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_table(table, widths=None, header=True, font_size=10):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if header and ri == 0:
                set_cell_shading(cell, "D9E2F3")
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run_font(run, size=font_size, bold=(header and ri == 0))
    if header:
        set_repeat_table_header(table.rows[0])


def add_data_table(doc, headers, rows, widths=None, font_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], text, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
    for values in rows:
        cells = table.add_row().cells
        for i, text in enumerate(values):
            set_cell_text(cells[i], str(text), size=font_size)
    format_table(table, widths=widths, header=True, font_size=font_size)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def remove_template_body_after_basic_table(doc):
    body = doc._element.body
    first_p = doc.paragraphs[0]._p
    first_table = doc.tables[0]._tbl
    for child in list(body):
        if child is first_p or child is first_table or child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def update_basic_table(doc):
    heading = doc.paragraphs[0]
    heading.text = "基本情况表"
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(8)
    set_run_font(heading.runs[0], chinese=HEADING_FONT_CN, latin="Arial", size=16, bold=True)

    values = [
        ("交底书编号", "（不填）"),
        ("交底书名称", PATENT_TITLE),
        ("拟申请的专利类型", "发明（拟申请）  实用新型□  外观设计□  同日申请发明实用新型□"),
        ("其他要求", "建议提前公开并同时请求实质审查；最终由专利代理人确认"),
        ("技术负责人（撰写人）及联系", "刘颖（联系电话、邮箱待补充）"),
        ("事务处理人及联系", "待填写"),
        ("发明人（全部）", "待确认（应包括对本方案作出创造性贡献的项目成员）"),
        ("第一发明人国籍及身份证号", "待填写"),
        ("权利人（单位）", "待确认"),
        ("关联项目", "DVCL异构图神经网络鲁棒性研究项目；项目类型及编号待填写"),
        ("关 键 字", "异构图神经网络；结构攻击；元路径；拓扑净化；特征诱导图；对比学习"),
        ("技术方向", "人工智能、图机器学习、数据安全与鲁棒学习"),
    ]
    table = doc.tables[0]
    while len(table.rows) < len(values):
        table.add_row()
    while len(table.rows) > len(values):
        table._tbl.remove(table.rows[-1]._tr)
    for row, (key, value) in zip(table.rows, values):
        set_cell_text(row.cells[0], key, bold=True, size=10)
        set_cell_shading(row.cells[0], "E7E6E6")
        set_cell_text(row.cells[1], value, size=10)
    format_table(table, widths=[4.1, 12.0], header=False, font_size=10)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = ""
    run = paragraph.add_run(f"{SHORT_TITLE}    第 ")
    set_run_font(run, size=9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9)


def reset_footer(footer):
    # The source template stores its original page number inside a content
    # control, which python-docx does not expose through footer.paragraphs.
    for child in list(footer._element):
        footer._element.remove(child)
    add_page_number(footer.add_paragraph())


def get_font(size, bold=False):
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf" if bold else r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_text(draw, text, font, max_width):
    lines = []
    current = ""
    for char in text:
        candidate = current + char
        width = draw.textbbox((0, 0), candidate, font=font)[2]
        if width <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def draw_centered_text(draw, box, text, font, fill="black", line_gap=8):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, max(20, x2 - x1 - 24))
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total = sum(heights) + line_gap * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    for line, h in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += h + line_gap


def draw_box(draw, box, text, font, width=4, fill="#F5F5F5"):
    draw.rounded_rectangle(box, radius=12, outline="black", width=width, fill=fill)
    draw_centered_text(draw, box, text, font)


def draw_arrow(draw, start, end, width=5, head=18):
    draw.line([start, end], fill="black", width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    p1 = (
        end[0] - head * math.cos(angle - math.pi / 6),
        end[1] - head * math.sin(angle - math.pi / 6),
    )
    p2 = (
        end[0] - head * math.cos(angle + math.pi / 6),
        end[1] - head * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, p1, p2], fill="black")


def new_canvas(title):
    image = Image.new("RGB", (1800, 1080), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(48, bold=True)
    bbox = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((1800 - bbox[2]) / 2, 35), title, fill="black", font=title_font)
    return image, draw


def build_figures(output_dir):
    output_dir.mkdir(parents=True, exist_ok=True)
    font = get_font(34)
    small = get_font(29)

    image, draw = new_canvas("DVCL总体结构")
    boxes = {
        "input": (60, 420, 310, 620),
        "topo": (430, 190, 780, 390),
        "feat": (430, 680, 780, 880),
        "enc1": (910, 190, 1210, 390),
        "enc2": (910, 680, 1210, 880),
        "align": (1280, 390, 1570, 680),
        "out": (1600, 420, 1760, 620),
    }
    draw_box(draw, boxes["input"], "异构图、节点特征及训练标签", font)
    draw_box(draw, boxes["topo"], "元路径诱导、边置信度估计、语义加权与两阶段过滤", small)
    draw_box(draw, boxes["feat"], "特征归一化、相似度计算与K近邻建图", small)
    draw_box(draw, boxes["enc1"], "第一图编码器\n拓扑表示", font)
    draw_box(draw, boxes["enc2"], "第二图编码器\n特征表示", font)
    draw_box(draw, boxes["align"], "同节点跨视图对齐\n表示融合\n联合损失优化", small)
    draw_box(draw, boxes["out"], "节点类别与置信度", small)
    draw_arrow(draw, (310, 470), (430, 290))
    draw_arrow(draw, (310, 570), (430, 780))
    draw_arrow(draw, (780, 290), (910, 290))
    draw_arrow(draw, (780, 780), (910, 780))
    draw_arrow(draw, (1210, 290), (1340, 420))
    draw_arrow(draw, (1210, 780), (1340, 650))
    draw_arrow(draw, (1570, 535), (1600, 535))
    draw.line([(1060, 405), (1060, 665)], fill="black", width=4)
    draw_arrow(draw, (1060, 535), (1280, 535), width=4)
    p1 = output_dir / "figure1_overall.png"
    image.save(p1, dpi=(240, 240))

    image, draw = new_canvas("净化拓扑视图构造流程")
    row1 = [
        ((50, 180, 300, 340), "关系类型邻接矩阵"),
        ((380, 180, 680, 340), "按元路径进行矩阵连乘"),
        ((760, 180, 1070, 340), "节点特征校准边分数"),
        ((1150, 180, 1450, 340), "元路径级阈值过滤"),
    ]
    for box, text in row1:
        draw_box(draw, box, text, small)
    for a, b in zip(row1, row1[1:]):
        draw_arrow(draw, (a[0][2], 260), (b[0][0], 260))
    draw_box(draw, (600, 520, 930, 710), "各元路径图的图注意力编码", small)
    draw_box(draw, (1010, 520, 1310, 710), "语义注意力权重", font)
    draw_box(draw, (1390, 520, 1740, 710), "加权合并并执行全局阈值过滤", small)
    draw_arrow(draw, (1300, 340), (820, 520))
    draw_arrow(draw, (930, 615), (1010, 615))
    draw_arrow(draw, (1310, 615), (1390, 615))
    draw_box(draw, (680, 850, 1120, 1010), "输出净化拓扑视图 G_topo", font)
    draw_arrow(draw, (1565, 710), (1120, 925))
    p2 = output_dir / "figure2_topology.png"
    image.save(p2, dpi=(240, 240))

    image, draw = new_canvas("特征诱导视图构造流程")
    steps = [
        ((60, 410, 330, 620), "目标类型节点特征矩阵 X"),
        ((420, 410, 690, 620), "L2归一化"),
        ((780, 410, 1080, 620), "计算两两余弦相似度"),
        ((1170, 410, 1450, 620), "为每个节点选取Top-k邻居"),
        ((1530, 410, 1760, 620), "输出特征视图 G_feat"),
    ]
    for box, text in steps:
        draw_box(draw, box, text, small)
    for a, b in zip(steps, steps[1:]):
        draw_arrow(draw, (a[0][2], 515), (b[0][0], 515))
    draw_centered_text(draw, (1150, 680, 1470, 810), "可选：有向、对称或互为近邻模式", small)
    draw_centered_text(draw, (570, 730, 1180, 930), "该视图不以观测异构图中的候选边作为建图依据，\n在结构边受到扰动时提供互补参照", font)
    p3 = output_dir / "figure3_feature.png"
    image.save(p3, dpi=(240, 240))

    image, draw = new_canvas("训练与推理处理流程")
    coords = [
        ((70, 180, 390, 340), "S1 获取异构图并确定目标节点类型"),
        ((520, 180, 840, 340), "S2 构造净化拓扑视图"),
        ((970, 180, 1290, 340), "S3 构造特征诱导视图"),
        ((1410, 180, 1730, 340), "S4 双编码器获得节点表示"),
        ((1410, 650, 1730, 830), "S5 融合表示并输出节点分类结果"),
        ((970, 650, 1290, 830), "S6 计算监督损失与跨视图对比损失"),
        ((520, 650, 840, 830), "S7 根据验证指标保存最佳参数"),
        ((70, 650, 390, 830), "S8 恢复最佳参数执行推理"),
    ]
    for box, text in coords:
        draw_box(draw, box, text, small)
    for idx in range(3):
        draw_arrow(draw, (coords[idx][0][2], 260), (coords[idx + 1][0][0], 260))
    draw_arrow(draw, (1570, 340), (1570, 650))
    for idx in range(4, 7):
        draw_arrow(draw, (coords[idx][0][0], 740), (coords[idx + 1][0][2], 740))
    draw_arrow(draw, (680, 650), (680, 500))
    draw_arrow(draw, (680, 500), (1570, 500))
    draw_centered_text(draw, (760, 430, 1470, 560), "训练阶段循环更新；推理阶段仅执行视图构造、编码、融合与分类", small)
    p4 = output_dir / "figure4_flow.png"
    image.save(p4, dpi=(240, 240))
    return [p1, p2, p3, p4]


def add_figure(doc, path, caption, width_cm=15.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_with_next = False
    cap.paragraph_format.keep_together = True
    r = cap.add_run(caption)
    set_run_font(r, size=10.5)


def build_document(template_path, output_path):
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(template_path))
    remove_template_body_after_basic_table(doc)
    update_basic_table(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.4)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            reset_footer(footer)

    with tempfile.TemporaryDirectory(prefix="dvcl_patent_figures_") as temp_dir:
        figures = build_figures(Path(temp_dir))

        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        add_center_title(doc, "专利申请交底书", 20)
        add_heading(doc, "发明名称", 1)
        add_body(doc, PATENT_TITLE, first_line=False)

        add_heading(doc, "所属技术领域", 1)
        add_body(
            doc,
            "本发明涉及人工智能、图机器学习和数据安全技术领域，具体涉及一种在异构图的边关系可能受到恶意插入、删除或其他结构扰动时，利用净化语义拓扑视图与特征诱导视图进行双视图表示学习、跨视图对比约束和节点分类的方法、系统、电子设备及计算机可读存储介质。",
        )

        add_heading(doc, "现有技术", 1)
        add_heading(doc, "1. 主要技术", 2)
        add_body(
            doc,
            "异构图用于描述包含多种节点类型和多种关系类型的复杂数据。例如，在学术网络中可同时包含论文、作者、会议或研究主题等节点，并包含撰写、发表、引用或隶属等关系；在推荐网络、社交网络和生物医学网络中也存在类似的多实体、多关系结构。异构图节点分类的目标，是利用节点属性和关系结构预测目标类型节点的类别。",
        )
        add_body(
            doc,
            "现有异构图神经网络通常使用元路径表达复合语义关系。元路径是由节点类型和关系类型组成的关系序列，通过沿元路径聚合多跳邻居，可将原始异构图投影为面向目标节点类型的一个或多个同构语义图。随后，图卷积、图注意力或图Transformer等编码器在语义图上进行消息传播，并通过语义注意力融合不同元路径的节点表示。",
        )
        add_body(
            doc,
            "图结构攻击通过添加、删除或重连少量边，改变节点局部邻域和消息传播路径。在异构图中，一条底层关系边可能同时参与多条元路径组合，因而少量扰动可被高阶语义传播放大，使大量目标节点的元路径邻域发生改变。受污染邻居的信息被编码器持续聚合后，会导致节点表示和分类结果偏离正常状态。",
        )
        add_heading(doc, "2. 主要方法", 2)
        add_body(
            doc,
            "第一类方法在观测拓扑上直接训练异构图神经网络，利用关系类型、元路径和语义注意力提高表示能力。该类方法主要关注干净数据上的语义建模，通常没有显式判断候选边是否可靠。",
        )
        add_body(
            doc,
            "第二类方法对观测图进行边剪枝、相似性筛选、概率估计或图结构学习。例如，根据相邻节点属性相似度计算边置信度，删除低置信边；或者为不同元路径学习权重，降低异常语义通道的影响。这些方法能够一定程度上抑制结构噪声。",
        )
        add_body(
            doc,
            "第三类方法通过数据增强、自监督学习或对比学习构造多个图视图，并要求同一节点在不同增强视图中的表示保持一致。常见增强仍然由当前观测拓扑进行删边、加边、子图采样或特征遮蔽，因此不同视图可能共享同一结构污染源。",
        )
        add_heading(doc, "3. 现有方法的不足", 2)
        add_numbered(doc, "（1）", "防御过程仍高度依赖受攻击拓扑。即使对边进行打分或剪枝，候选邻域仍来自被扰动后的异构图；当恶意边在特征上具有伪装性或结构污染较重时，仅依靠拓扑证据难以恢复稳定表示。")
        add_numbered(doc, "（2）", "元路径级净化与语义级融合缺少协同。仅设置统一阈值容易忽略不同元路径的密度和语义差异；仅依赖语义注意力又可能给已经受污染的元路径分配较高权重。")
        add_numbered(doc, "（3）", "拓扑信号和节点属性信号之间缺少节点级一致性约束。简单拼接两类特征不能保证同一节点在两个信息来源下具有一致语义，也不能主动分离不同节点的表示。")
        add_numbered(doc, "（4）", "传统单分支模型在结构攻击下缺少可替代的信息通路。拓扑邻域一旦被污染，分类器没有相对独立的参照视图用于校正预测。")
        add_numbered(doc, "（5）", "部分多视图方案只在输入层进行随机增强，未同时解决异构元路径语义保留、可疑边过滤以及跨视图表示对齐问题。")

        add_heading(doc, "发明创造的目的", 1)
        add_body(
            doc,
            "本发明的目的在于提供一种基于双视图对比学习的鲁棒异构图节点分类方案，在保留异构关系语义的同时，降低节点分类结果对受扰动图拓扑的单一依赖。为此，本发明从同一批目标节点构造净化拓扑视图和特征诱导视图，并对两个视图的节点表示进行联合分类和对称跨视图对齐。",
        )
        add_numbered(doc, "（1）", "通过元路径级边置信度过滤与语义加权后的全局过滤，分层去除低可信语义连接，减少结构扰动沿元路径传播和放大的影响。")
        add_numbered(doc, "（2）", "通过目标节点属性构造特征诱导近邻图，为受攻击拓扑提供相对独立的结构参照，使模型在部分拓扑边不可信时仍能够获得稳定邻域信息。")
        add_numbered(doc, "（3）", "通过同一节点在两个视图中的正样本配对以及不同节点之间的负样本区分，建立节点级跨视图一致性，降低两个视图的语义偏差。")
        add_numbered(doc, "（4）", "通过监督分类损失、辅助语义分支损失和跨视图对比损失联合优化，实现从语义拓扑净化、双视图编码到节点分类的端到端训练。")
        add_numbered(doc, "（5）", "提供可替换的近邻建图、图编码和视图融合方式，使该方案能够用于学术网络、推荐网络、社交网络、生物医学网络等不同异构图场景。")

        add_heading(doc, "发明创造的技术方案", 1)
        add_heading(doc, "1. 总体方案", 2)
        add_body(
            doc,
            "本发明将面向目标节点类型的鲁棒异构图节点分类过程划分为数据输入、净化拓扑视图构造、特征诱导视图构造、双视图编码、表示融合、跨视图对比学习和分类输出七个功能阶段。两个视图包含相同的目标节点集合，但边集合由不同证据生成：净化拓扑视图保留元路径语义，特征诱导视图主要由节点属性相似性生成。总体结构如图1所示。",
        )
        add_figure(doc, figures[0], "图1  DVCL总体结构", width_cm=14.0)
        add_body(
            doc,
            "给定异构图G=(V,E,A,R)，其中V为节点集合，E为边集合，A为节点类型集合，R为关系类型集合。设V_t={v_1,…,v_n}为待分类的目标类型节点集合，X∈R^(n×d)为目标节点特征矩阵，Y为已知标签集合。训练阶段仅使用训练节点标签，推理阶段输出目标节点类别及其预测置信度。",
        )

        add_heading(doc, "2. 关键符号", 2)
        add_data_table(
            doc,
            ["符号", "含义", "符号", "含义"],
            [
                ("G", "输入异构图", "V_t", "目标类型节点集合"),
                ("P_m", "第m条元路径", "A^(m)", "第m条元路径诱导邻接矩阵"),
                ("w_ij^(m)", "元路径m下边(i,j)的置信度", "β_m", "元路径m的语义注意力权重"),
                ("G_topo", "净化拓扑视图", "G_feat", "特征诱导视图"),
                ("Z_topo", "拓扑视图节点表示", "Z_feat", "特征视图节点表示"),
                ("γ_m", "元路径级过滤阈值", "γ_s", "融合语义图的全局过滤阈值"),
                ("k", "特征近邻数", "τ", "跨视图对比学习温度参数"),
            ],
            widths=[2.2, 5.8, 2.2, 5.8],
            font_size=9.5,
        )

        add_heading(doc, "3. 净化拓扑视图构造", 2)
        add_heading(doc, "3.1 元路径诱导图构造", 3)
        add_body(
            doc,
            "根据业务语义或数据模式确定元路径集合P={P_1,…,P_M}。对于元路径P_m=(r_1,r_2,…,r_L)，获取各关系类型的邻接矩阵，对每一关系邻接矩阵进行行归一化，并按照元路径顺序进行矩阵连乘，得到以目标节点为起点和终点的元路径诱导邻接矩阵：",
        )
        add_formula(doc, "A^(m) = Ā_(r1) Ā_(r2) … Ā_(rL)", 1)
        add_body(doc, "式中，Ā_(rl)表示关系r_l对应的行归一化邻接矩阵。矩阵元素A_ij^(m)表示节点v_i与节点v_j沿元路径P_m的可达强度或连接权重。")

        add_heading(doc, "3.2 特征校准的边置信度估计", 3)
        add_body(
            doc,
            "仅根据元路径可达性无法判断连接是否符合节点属性语义，因此对每一条候选边(i,j)使用目标节点特征进行校准。在一个具体实施方式中，先计算结构强度与节点特征内积的乘积：",
        )
        add_formula(doc, "s_ij^(m) = A_ij^(m) · x_i^T x_j", 2)
        add_body(
            doc,
            "再对边分数执行度归一化、行归一化、对称归一化或其他尺度标准化。在对称度归一化的一个实施方式中，可采用：",
        )
        add_formula(doc, "w_ij^(m) = (d_i^(m))^(-1) s_ij^(m) (d_j^(m))^(-1)", 3)
        add_body(doc, "式中，d_i^(m)=Σ_j s_ij^(m)。工程实现时可在分母加入预设极小正数，以避免零度节点造成数值异常。所得w_ij^(m)作为元路径m下候选边(i,j)的置信度。")

        add_heading(doc, "3.3 元路径级第一阶段过滤", 3)
        add_body(
            doc,
            "针对不同元路径的结构密度、数值范围和攻击敏感性分别设置阈值γ_m，保留置信度不低于阈值的候选边：",
        )
        add_formula(doc, "E_pur^(m) = {(i,j) | w_ij^(m) ≥ γ_m}", 4)
        add_body(doc, "阈值γ_m可以是固定值、基于验证集选择的值、分位数阈值、按节点保留固定比例的阈值，或者由可学习门控模块产生的阈值。由此得到多个净化元路径图。")

        add_heading(doc, "3.4 语义注意力与第二阶段过滤", 3)
        add_body(
            doc,
            "针对每一个净化元路径图，使用图编码器得到元路径特定节点表示h_i^(m)，再通过语义注意力网络计算该元路径对当前任务的重要性。在一个实施方式中，元路径语义分数和归一化权重分别为：",
        )
        add_formula(doc, "e_m = (1/n) Σ_i q^T tanh(W_s h_i^(m) + b_s)", 5)
        add_formula(doc, "β_m = exp(e_m) / Σ_(m'=1)^M exp(e_m')", 6)
        add_body(
            doc,
            "利用语义权重对来自不同元路径的边置信度进行加权合并；对于同一节点对在多条元路径中重复出现的情况，将其对应权重累加或按预设聚合函数合并：",
        )
        add_formula(doc, "w̄_ij = Σ_(m=1)^M β_m w_ij^(m)", 7)
        add_body(doc, "对融合后的边分数应用全局语义阈值γ_s，得到净化拓扑视图：")
        add_formula(doc, "E_topo = {(i,j) | w̄_ij ≥ γ_s}，G_topo=(V_t,E_topo)", 8)
        add_body(
            doc,
            "第一阶段过滤解决各元路径内部的低置信连接问题，第二阶段过滤结合任务相关的语义权重处理跨元路径融合后的连接问题。该分层方式避免单一阈值同时承担元路径内部去噪和跨语义融合两种功能。具体流程如图2所示。",
        )
        add_figure(doc, figures[1], "图2  净化拓扑视图构造流程")

        add_heading(doc, "4. 特征诱导视图构造", 2)
        add_body(
            doc,
            "为在结构边受到攻击时提供互补信息，本发明从目标节点属性构造第二视图。优选地，该建图过程不使用当前观测异构拓扑中的候选边。对于目标节点特征x_i，先执行L2归一化：",
        )
        add_formula(doc, "x̃_i = x_i / ||x_i||_2", 9)
        add_body(doc, "随后计算任意两个目标节点之间的余弦相似度：")
        add_formula(doc, "sim(i,j) = x̃_i^T x̃_j", 10)
        add_body(doc, "对于每个节点v_i，选择除自身外相似度最高的k个节点构成近邻集合：")
        add_formula(doc, "N_k(i) = TopK_(j≠i) sim(i,j)", 11)
        add_body(doc, "根据近邻集合建立特征诱导边集合和特征诱导图：")
        add_formula(doc, "E_feat = {(i,j) | j∈N_k(i)}，G_feat=(V_t,E_feat)", 12)
        add_body(
            doc,
            "近邻图可采用有向模式、对称模式或互为近邻模式。有向模式保留每个节点指向其Top-k节点的边；对称模式将任一方向出现的近邻关系转换为双向边；互为近邻模式仅保留两个节点彼此进入对方Top-k集合的边。对于大规模图，可使用分块相似度计算或近似最近邻索引，避免显式存储完整的n×n相似度矩阵。",
        )
        add_body(
            doc,
            "节点特征可以是原始属性，也可以是在不改变本发明双视图构思的前提下，由可信侧信息、预训练特征、时间上早于攻击的快照特征或关系分组的统计特征形成的增强属性。优选实施方式直接采用原始属性，从而降低第二视图与受攻击拓扑的耦合。流程如图3所示。",
        )
        add_figure(doc, figures[2], "图3  特征诱导视图构造流程")

        add_heading(doc, "5. 双视图编码与融合分类", 2)
        add_heading(doc, "5.1 独立图编码", 3)
        add_body(
            doc,
            "为净化拓扑视图和特征诱导视图分别设置参数独立的第一图编码器f_topo和第二图编码器f_feat。优选地，两个编码器采用包含自环的图注意力网络；也可以采用图卷积网络、GraphSAGE、图Transformer或其他能够根据邻接关系传播节点信息的图编码器。",
        )
        add_formula(doc, "Z_topo = f_topo(X,G_topo)", 13)
        add_formula(doc, "Z_feat = f_feat(Dropout(X),G_feat)", 14)
        add_body(
            doc,
            "式中，Z_topo和Z_feat具有相同的节点顺序。对特征视图输入执行随机遮蔽或dropout，可减少编码器对少量属性维度的过拟合；训练结束后的推理阶段关闭随机遮蔽。两个编码器独立设置参数，可避免将两个视图强制映射为完全相同的消息传播过程。",
        )

        add_heading(doc, "5.2 表示融合与分类", 3)
        add_body(doc, "对于节点v_i，分别得到拓扑视图表示z_i^topo和特征视图表示z_i^feat。优选地，对两个表示进行拼接：")
        add_formula(doc, "z_i = z_i^topo || z_i^feat", 15)
        add_body(doc, "将融合表示输入线性分类器或多层感知机，获得节点类别预测：")
        add_formula(doc, "o_i = W_c z_i + b_c", 16)
        add_body(
            doc,
            "除直接拼接外，还可根据两个表示计算节点相关门控系数g_i，并采用g_i z_i^topo+(1-g_i)z_i^feat进行加权求和，或者采用[g_i z_i^topo || (1-g_i)z_i^feat]进行门控拼接。该替代方案允许模型根据节点状态调整两个视图的贡献。",
        )

        add_heading(doc, "6. 跨视图对比学习与联合训练", 2)
        add_heading(doc, "6.1 监督分类损失", 3)
        add_body(doc, "在有标签训练节点集合V_tr上计算交叉熵分类损失：")
        add_formula(doc, "L_cls = -(1/|V_tr|) Σ_(i∈V_tr) log[exp(o_(i,y_i))/Σ_c exp(o_(i,c))]", 17)

        add_heading(doc, "6.2 对称跨视图对比损失", 3)
        add_body(
            doc,
            "先对两个视图的节点表示执行L2归一化。同一目标节点在两个视图中的表示构成正样本对，不同目标节点在相反视图中的表示构成负样本。以拓扑视图表示为锚点、特征视图表示为对比对象时，对比损失为：",
        )
        add_formula(doc, "L_(t→f) = -(1/|V_tr|) Σ_i log{exp[(z̃_i^topo)^T z̃_i^feat/τ] / Σ_j exp[(z̃_i^topo)^T z̃_j^feat/τ]}", 18)
        add_body(doc, "交换锚点和对比对象，得到特征视图到拓扑视图方向的损失L_(f→t)。最终对称跨视图损失为：")
        add_formula(doc, "L_cl = 1/2 [L_(t→f) + L_(f→t)]", 19)
        add_body(
            doc,
            "上述对称约束同时要求拓扑表示能够识别对应的特征表示，并要求特征表示能够识别对应的拓扑表示。相比单向约束，能够减少某一视图单方面追随另一视图而产生的表示偏置。",
        )

        add_heading(doc, "6.3 辅助语义监督和总损失", 3)
        add_body(
            doc,
            "为使语义注意力模块直接接收分类监督，可在净化元路径图的语义融合表示上设置辅助分类头，计算辅助语义分支损失L_HAN。基本实施方式的总训练目标为：",
        )
        add_formula(doc, "L = λ_HAN L_HAN + L_cls + λ_DVCL L_cl", 20)
        add_body(
            doc,
            "其中λ_HAN和λ_DVCL分别控制辅助语义监督和跨视图对比学习的强度。语义注意力模块、两个图编码器和最终分类器通过同一优化器联合更新。根据需要，还可以加入单视图分类损失、类别原型对比损失或有监督对比损失，但这些附加损失不改变本发明的双视图构造与跨视图对齐主线。",
        )

        add_heading(doc, "7. 处理流程", 2)
        add_numbered(doc, "S1.", "获取异构图、节点类型、关系类型、目标节点特征及训练节点标签；确定待分类的目标节点类型和元路径集合。")
        add_numbered(doc, "S2.", "沿各元路径组合关系邻接矩阵，得到多个目标节点同构语义图；计算特征校准边置信度，并执行元路径级第一阶段过滤。")
        add_numbered(doc, "S3.", "在净化元路径图上学习语义注意力权重，对各元路径的边置信度进行加权合并，并执行全局第二阶段过滤，得到净化拓扑视图。")
        add_numbered(doc, "S4.", "对目标节点特征进行归一化和相似度计算，为每个节点选择k个近邻，得到特征诱导视图。")
        add_numbered(doc, "S5.", "利用两个独立图编码器分别编码净化拓扑视图和特征诱导视图，得到按节点对应的两组表示。")
        add_numbered(doc, "S6.", "融合两组表示并输出节点分类结果；在训练阶段计算监督分类损失、辅助语义分支损失和对称跨视图对比损失。")
        add_numbered(doc, "S7.", "根据总损失反向传播并更新模型参数；根据验证节点上的预设指标保存最佳参数组合。")
        add_numbered(doc, "S8.", "在推理阶段恢复最佳参数，重新构造或读取两个视图，输出目标节点类别及其置信度。")
        add_figure(doc, figures[3], "图4  训练与推理处理流程")

        add_heading(doc, "8. 系统、设备及存储介质方案", 2)
        add_body(doc, "对应上述方法，本发明还提供一种鲁棒异构图节点分类系统，包括以下模块：")
        add_numbered(doc, "（1）", "数据获取模块，用于获取异构图、目标节点特征、元路径定义和训练节点标签。")
        add_numbered(doc, "（2）", "元路径投影模块，用于沿元路径组合关系邻接矩阵，生成目标节点同构语义图。")
        add_numbered(doc, "（3）", "拓扑净化模块，用于计算特征校准边置信度，执行元路径级过滤、语义注意力加权以及全局过滤。")
        add_numbered(doc, "（4）", "特征建图模块，用于根据目标节点属性相似度生成特征诱导近邻图。")
        add_numbered(doc, "（5）", "双视图编码模块，用于通过第一图编码器和第二图编码器分别获得拓扑表示和特征表示。")
        add_numbered(doc, "（6）", "跨视图学习模块，用于构造同一节点的跨视图正样本对和不同节点的负样本，计算对称跨视图对比损失。")
        add_numbered(doc, "（7）", "分类与训练模块，用于融合两个视图的表示、输出节点类别，并根据联合损失更新模型参数。")
        add_body(
            doc,
            "本发明还提供一种电子设备，包括处理器和存储器，存储器中存储有计算机程序；处理器执行所述计算机程序时实现上述方法步骤。本发明还提供一种计算机可读存储介质，其上存储有计算机程序，所述程序被处理器执行时实现上述方法步骤。",
        )

        add_heading(doc, "相对现有技术的优点和效果", 1)
        add_numbered(doc, "（1）", "降低结构依赖。特征诱导视图不以当前观测异构图中的候选边作为主要建图依据，在恶意插边或删边发生时能够提供与拓扑视图来源不同的邻域证据。")
        add_numbered(doc, "（2）", "分层净化异构语义。元路径级阈值用于处理各语义通道内部的低可信连接，语义注意力加权后的全局阈值用于处理跨通道融合结果，从而兼顾元路径差异和整体任务相关性。")
        add_numbered(doc, "（3）", "建立节点级跨视图一致性。对称InfoNCE目标将同一节点在两个视图中的表示相互拉近，并区分不同节点，避免简单融合时两个视图语义不一致。")
        add_numbered(doc, "（4）", "保留异构语义表达能力。净化拓扑视图仍通过元路径和语义注意力表示多类型关系，不需要将原始异构图简单退化为不区分关系类型的同构图。")
        add_numbered(doc, "（5）", "端到端联合优化。拓扑语义分支、双视图编码器、对比目标和分类器可以统一训练，减少多个独立阶段分别优化造成的目标不一致。")
        add_numbered(doc, "（6）", "具有可扩展性。近邻模式、过滤方式、编码器和融合方式可以替换，并可通过分块相似度或近似近邻技术应用于较大规模的异构图。")
        add_body(
            doc,
            "需要说明的是，上述技术效果针对异构图结构边受到扰动的场景。若节点属性本身也遭受攻击，可进一步采用可信属性筛选、属性去噪、联合拓扑与特征攻击训练等措施；该情形不影响本发明所述双视图构造和跨视图对齐技术方案。",
        )

        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        add_heading(doc, "附图说明", 1)
        add_numbered(doc, "图1", "为本发明DVCL鲁棒异构图节点分类方法的总体结构示意图。")
        add_numbered(doc, "图2", "为本发明净化拓扑视图的构造流程示意图。")
        add_numbered(doc, "图3", "为本发明特征诱导视图的构造流程示意图。")
        add_numbered(doc, "图4", "为本发明训练与推理处理流程示意图。")
        add_body(doc, "附图中的模块边界用于说明信息流和处理顺序，并不限定软件进程、硬件单元或部署节点的物理边界。多个模块可以合并实现，一个模块也可以拆分为多个子模块实现。")

        add_heading(doc, "具体实施方式", 1)
        add_heading(doc, "实施例一：学术异构图中的论文节点分类", 2)
        add_body(
            doc,
            "本实施例以包含论文、作者和研究主题或会议等节点的学术异构图为例。目标节点类型为论文，目标任务为预测论文所属研究领域。输入包括论文特征矩阵、论文与作者之间的关系、论文与主题或会议之间的关系，以及部分论文的领域标签。",
        )
        add_numbered(doc, "（1）", "选择“论文-作者-论文”和“论文-主题-论文”作为两条元路径。分别对“论文到作者”“作者到论文”“论文到主题”和“主题到论文”的关系邻接矩阵进行行归一化，并按元路径进行矩阵连乘。")
        add_numbered(doc, "（2）", "对各元路径诱导矩阵中的非零元素计算节点特征校准边分数，并进行归一化。根据验证集分别选择两条元路径的过滤阈值，生成净化元路径图。")
        add_numbered(doc, "（3）", "在两个净化元路径图上使用图注意力编码器获得语义表示，计算两条元路径的语义注意力权重；将边置信度按语义权重合并后进行全局过滤，得到论文节点的净化拓扑视图。")
        add_numbered(doc, "（4）", "对论文属性向量执行L2归一化，计算论文间余弦相似度。一个具体配置中设置k=20，构造有向、对称或互为近邻的特征诱导图。")
        add_numbered(doc, "（5）", "两个视图分别使用带自环的多头图注意力编码器。一个可行配置中，每个视图的隐藏维度为128、注意力头数为4，特征视图的输入遮蔽率为0.2。上述数值仅为示例，不构成保护范围限制。")
        add_numbered(doc, "（6）", "将两个视图的节点表示拼接后输入分类器，并在训练节点上同时计算分类损失、辅助语义分支损失和对称跨视图对比损失。一个具体配置中λ_HAN=1、λ_DVCL=1，温度参数τ=0.5。")
        add_numbered(doc, "（7）", "采用Adam或其他梯度优化算法更新全部可训练参数；使用验证集损失或验证准确率选择最佳模型，并在测试节点或待预测节点上输出论文类别。")

        add_heading(doc, "实施例二：推荐、社交或生物医学异构图", 2)
        add_body(
            doc,
            "在推荐场景中，目标节点可以是用户或物品，关系类型可以包括点击、购买、收藏、属于类别等；在社交场景中，目标节点可以是账户，关系类型可以包括关注、转发、评论和共同群组；在生物医学场景中，目标节点可以是药物、疾病或蛋白质，关系类型可以包括相互作用、治疗、关联和表达。实施时根据目标任务定义首尾均为目标节点类型的元路径，并重复执行净化拓扑视图和特征诱导视图构造。",
        )
        add_body(
            doc,
            "当目标节点数量较多时，不直接形成完整相似度矩阵，而将目标节点划分为多个批次，分块计算相似度并仅保留每行Top-k结果；也可以使用局部敏感哈希、向量索引或其他近似最近邻算法生成候选集合，再计算精确相似度。该实现降低内存开销，但不改变由节点属性确定特征诱导邻居的技术实质。",
        )

        add_heading(doc, "实施例三：训练服务与在线推理部署", 2)
        add_body(
            doc,
            "系统可部署为离线训练服务和在线或批量推理服务。离线训练服务读取异构图快照，生成元路径矩阵、净化拓扑视图和特征诱导视图，联合训练语义注意力模块、两个图编码器和分类器，并将最佳参数、元路径定义、过滤阈值和近邻参数作为同一模型版本保存。",
        )
        add_body(
            doc,
            "推理服务加载对应模型版本。对于静态快照，可复用已保存的两个视图；对于新增节点或周期性更新的异构图，可重新计算受影响的元路径候选边和特征近邻，随后执行双视图编码和融合分类。输出可以包括预测类别、类别概率、两个视图的单独诊断分数以及门控融合权重，以支持下游风险分析。",
        )
        add_body(
            doc,
            "为了保证训练与推理一致性，优选将数据版本、目标节点顺序、元路径集合、节点特征处理方式、阈值、近邻参数和模型参数写入配置清单。对于安全敏感业务，还可对输入图版本、模型版本和预测结果进行审计记录。",
        )

        add_heading(doc, "算法流程示例", 2)
        add_data_table(
            doc,
            ["步骤", "输入", "主要处理", "输出"],
            [
                ("1", "异构图G、元路径集合P", "关系矩阵归一化并按元路径连乘", "A^(1)…A^(M)"),
                ("2", "元路径矩阵、节点特征X", "计算并归一化特征校准边分数", "w_ij^(m)"),
                ("3", "边置信度、阈值γ_m", "执行元路径级过滤", "净化元路径图"),
                ("4", "净化元路径图", "图编码与语义注意力加权", "β_m"),
                ("5", "β_m、w_ij^(m)、γ_s", "融合边分数并全局过滤", "G_topo"),
                ("6", "节点特征X、近邻数k", "归一化、余弦相似度与Top-k", "G_feat"),
                ("7", "G_topo、G_feat、X", "独立编码、表示融合与分类", "节点logits"),
                ("8", "logits、训练标签、两视图表示", "计算联合损失并更新参数", "训练后的DVCL模型"),
            ],
            widths=[1.4, 4.0, 6.6, 4.0],
            font_size=9.2,
        )

        add_heading(doc, "可替换实施方式", 1)
        add_numbered(doc, "（1）", "边置信度中的特征一致性项可以采用内积、余弦相似度、径向基函数、可学习相似度网络或基于类别原型的相似度。")
        add_numbered(doc, "（2）", "过滤条件可以采用固定阈值、分位数、Top-k、按节点保留比例、统计显著性检验或可学习软门控；第一阶段和第二阶段可以采用不同类型的过滤规则。")
        add_numbered(doc, "（3）", "语义权重可以通过全局语义注意力、节点相关语义注意力、关系门控网络或基于验证目标的自适应权重获得。")
        add_numbered(doc, "（4）", "特征诱导图可以采用K近邻图、互为近邻图、半径邻域图、稀疏相似度图或基于聚类中心的二部图，只要其边主要由节点属性或可信侧信息确定。")
        add_numbered(doc, "（5）", "双视图编码器可以采用相同网络结构但参数独立，也可以采用不同网络结构；图注意力编码器可替换为图卷积、采样聚合、图Transformer或其他消息传递网络。")
        add_numbered(doc, "（6）", "两个视图可采用拼接、加权求和、门控拼接、交叉注意力或多层感知机融合。")
        add_numbered(doc, "（7）", "跨视图对比学习可以使用全批次负样本、小批次负样本、内存队列、类别原型或有监督同类正样本；可以使用InfoNCE、温度缩放交叉熵或其他具有同节点对齐和异节点区分作用的对比目标。")
        add_numbered(doc, "（8）", "本发明既可以用于结构扰动后的防御推理，也可以在包含干净图与扰动图的训练集合上进行鲁棒训练；还可以与属性去噪或联合攻击训练组合。")

        add_heading(doc, "软件开发与运行环境", 1)
        add_body(
            doc,
            "一个具体软件实现可使用Python和深度学习框架完成，并使用异构图处理库管理多类型节点和关系。主要软件组件包括数据加载器、元路径变换器、边置信度计算器、图过滤器、特征近邻构造器、双视图模型、训练器、检查点管理器和评估器。实现语言、框架和硬件平台不构成对本发明保护范围的限制。",
        )
        add_data_table(
            doc,
            ["组件", "主要职责", "可选实现"],
            [
                ("数据加载器", "读取异构图、特征、标签和数据划分", "文件、数据库或图存储"),
                ("元路径变换器", "构造元路径诱导邻接关系", "稀疏矩阵乘法或逐路径遍历"),
                ("拓扑净化器", "边置信度、两阶段过滤与语义融合", "CPU稀疏计算或GPU张量计算"),
                ("特征建图器", "相似度计算与近邻选择", "精确Top-k或近似近邻索引"),
                ("双视图模型", "独立编码、融合、分类和对比损失", "图注意力或其他图编码网络"),
                ("训练与检查点", "联合优化、早停、参数保存与恢复", "单机、多GPU或分布式训练"),
            ],
            widths=[3.0, 7.0, 6.0],
            font_size=9.3,
        )

        add_heading(doc, "验证方案", 1)
        add_body(
            doc,
            "为验证本发明的技术效果，可在至少一个包含多种节点和关系类型的数据集上设置干净图、全局结构攻击图和目标结构攻击图。全局攻击按总边数的一定比例进行插边或删边；目标攻击为选定目标节点设置不同扰动预算。所有对比方法使用相同的数据划分、攻击后图、目标节点集合和随机种子。",
        )
        add_body(
            doc,
            "主要验证指标可以采用节点分类准确率及其多随机种子均值与标准差。还应设置以下消融对照：仅使用净化拓扑视图、仅使用特征诱导视图、使用双视图但移除跨视图对比损失、移除第一阶段或第二阶段过滤、改变近邻数k、改变λ_DVCL以及改变视图融合方式。由此分别验证双视图互补性、分层过滤和跨视图对齐的贡献。",
        )

        add_heading(doc, "要保护的发明内容的技术关键点", 1)
        add_numbered(doc, "1.", "一种面向异构图节点分类的双视图鲁棒学习总体方案：针对同一目标节点集合，构造净化拓扑视图和主要由节点属性确定的特征诱导视图，并基于两视图联合输出节点分类结果。")
        add_numbered(doc, "2.", "一种异构语义拓扑的两阶段净化机制：先在各元路径内部根据特征校准边置信度执行元路径级过滤，再根据语义注意力对各元路径边分数进行加权合并并执行全局过滤。")
        add_numbered(doc, "3.", "一种用于结构攻击防御的特征诱导建图机制：根据目标节点属性相似度建立近邻图，使第二视图不依赖或较少依赖当前受扰动拓扑中的候选边。")
        add_numbered(doc, "4.", "一种对称跨视图节点级对比机制：将同一节点在拓扑视图和特征视图中的表示作为双向正样本对，并将不同节点的跨视图表示作为负样本进行区分。")
        add_numbered(doc, "5.", "一种监督分类、辅助语义监督和跨视图对比约束的联合训练机制，使语义注意力模块、双视图编码器和分类器端到端协同优化。")
        add_numbered(doc, "6.", "与上述方法对应的系统模块、电子设备和计算机可读存储介质，以及近邻模式、过滤规则、图编码器和融合方式的等同替换方案。")

        add_heading(doc, "撰写权利要求时的建议保护层级", 1)
        add_body(
            doc,
            "建议独立方法权利要求至少包含：获取异构图和目标节点属性；根据元路径和边置信度构造净化拓扑视图；根据节点属性相似度构造特征诱导视图；分别编码两个视图；通过跨视图对比约束和监督分类目标训练；融合表示并输出节点类别。避免在独立权利要求中限定具体数据集、固定阈值、固定近邻数、GAT层数或具体损失权重。",
        )
        add_body(
            doc,
            "建议从属权利要求依次限定：元路径矩阵连乘；特征校准边置信度；元路径级阈值；语义注意力权重；全局第二阶段过滤；有向、对称或互为近邻建图；带自环的独立图编码器；特征遮蔽；拼接或门控融合；对称InfoNCE；辅助语义分类损失；最佳检查点恢复。另设置系统、电子设备和存储介质独立权利要求。",
        )

    doc.core_properties.title = PATENT_TITLE
    doc.core_properties.subject = "DVCL鲁棒异构图学习专利交底书"
    doc.core_properties.author = "刘颖"
    doc.core_properties.keywords = "DVCL, 异构图, 鲁棒学习, 元路径, 拓扑净化, 特征诱导图, 对比学习"
    doc.core_properties.comments = "根据DVCL论文与项目实现整理；联系人、发明人和权利人信息待确认。"
    doc.save(str(output_path))


def parse_args():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("template", type=Path, help="Source patent disclosure DOCX template")
    parser.add_argument("output", type=Path, help="Destination DOCX path")
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build_document(args.template.resolve(), args.output.resolve())
    print(args.output.resolve())
